import pandas as pd
from tqdm import tqdm
import numpy as np
import os
import openai
import time 
from urllib.request import urlopen
from bs4 import BeautifulSoup
import requests
import json

import os
import datetime
import pandas as pd
from email.mime.multipart import MIMEMultipart
from email.mime.application import MIMEApplication
from email import encoders
from email.mime.base import MIMEBase
from email.mime.text import MIMEText
from email.mime.image import MIMEImage
import smtplib
from datetime import datetime, timedelta
import time
import glob
import ast
import google.generativeai as genai
import utils.config as config
import docx

def run_genai(prompt):
    genai.configure(api_key=config.api_key)
    model = genai.GenerativeModel(config.model)
    response = model.generate_content(f'{prompt}', generation_config=genai.types.GenerationConfig(temperature=0.0, max_output_tokens=20000))
    try:
        return response.text
    except:
        return 'error response'

def parse_html(url):
    
    foa_title = url.split('files/')[1].split('.html')[0]
    
    html = urlopen(url).read()
    soup = BeautifulSoup(html, features="html.parser")

    for script in soup(["script", "style"]):
        script.extract()    # rip it out
    
    text = soup.get_text()
    
    lines = (line.strip() for line in text.splitlines())
    
    chunks = (phrase.strip() for line in lines for phrase in line.split("  "))
    
    text = '\n'.join(chunk for chunk in chunks if chunk)
    
    return text, foa_title

def parse_html_template(url):

    
    html = urlopen(url).read()
    soup = BeautifulSoup(html, features="html.parser")

    for script in soup(["script", "style"]):
        script.extract()    # rip it out
    
    text = soup.get_text()
    
    lines = (line.strip() for line in text.splitlines())
    
    chunks = (phrase.strip() for line in lines for phrase in line.split("  "))
    
    text = '\n'.join(chunk for chunk in chunks if chunk)
    
    return text

def extract_due_dates():
    standard_due_dates_url = 'https://grants.nih.gov/grants/how-to-apply-application-guide/due-dates-and-submission-policies/due-dates.htm'
    standard_due_dates_txt = parse_html_template(standard_due_dates_url)
    
    return standard_due_dates_txt
    
def extract_foa_template():
    foa_template_url = 'https://grants.nih.gov/grants/how-to-apply-application-guide/due-dates-and-submission-policies/due-dates.htm'
    foa_template = parse_html_template(foa_template_url)
    
    return foa_template

def generate_prompt(foa, current_datestr):
    standard_due_dates_txt = extract_due_dates()
    foa_template = extract_foa_template()
    prompt = f"""
    Using the the FOA, which is delimited by double back ticks, extract the following information from these 29 questions:

    1. Participating Organizations for FOA
    
    2. Components of Participating Organizations for FOA
    
    3. What is the Funding Opportunity Title for FOA
    
    4. What are the activity code(s) for FOA
    
    5. Announcement Type for FOA
    
    6. When is application due for FOA.
    
    7. What is the project start date for FOA
    
    8. What is the expiration date for FOA
    
    9. When is Letter of Intent due for FOA
    
    11. What is the award budget for FOA.  
    
    12. Are there any budget limitations for FOA
    
    13. What is the award  project period for FOA
    
    15. Is Clinical Trial allowed on FOA
    
    16. Is Cost Sharing required on FOA
    
    17. Is Cost Sharing recommended on FOA
    
    18. Are there any Page Limitations for FOA 
    
    19. What are the page limitations for FOA
    
    20. Are there any Limited submission requirements for FOA
    
    21. Can I submit more than one application for FOA
    
    22. Are Foreign Components allowed for FOA
    
    23. What are the other attachments for FOA
    
    25. Who are the agency contacts for FOA
    
    26. What are the section headers and subsections of the FOA
    
    27. What are other submission requirements that start with the phrase "specific to this NOFO."
    
    28. What are the key dates indicated in the table? Extract information from the table under key dates and specifically extract information from the column with the header that contains that the word "AIDS"
    
    29. What are the start dates indicated in the table? Extract information from the table under key dates and specifically extract information from the column with the header that contains that the word "earliest start date". Please select the first date after {current_datestr}.
    
    
    FOA:
    ``{foa}`` 
    
    Please make sure to enclose the information pertaining to the exact section where the answered was retrieved in square brackets and start with the phrase "Information found from".    
    Please do not include phrases like "here is the extracted information" in your output. 
    
    In your output, use this format:
    Question Number: Question
    
    Answer: <place your answer here>
    
    Please make sure that you went through each step in the prompt mentioned above.
    """
    return prompt
  
    
def generate_main_df():
    dict_questions = {1: 'Participating Organizations for FOA',
                    2: 'Components of Participating Organizations for FOA',
                    3: 'What is the Funding Opportunity Title for FOA',
                    4: 'What are the activity code(s) for FOA',
                    5: 'Announcement Type for FOA',
                    6: 'When is application due for FOA',
                    7: 'What is the project start date for FOA',
                    8: 'What is the expiration date for FOA',
                    9: 'When is Letter of Intent due for FOA',
                    11: 'What is the award budget for FOA',
                    12: 'Are there any budget limitations for FOA',
                    13: 'What is the award  project period for FOA',
                    15: 'Is Clinical Trial allowed on FOA',
                    16: 'Is Cost Sharing required on FOA',
                    17: 'Is Cost Sharing recommended on FOA',
                    18: 'Are there any Page Limitations for FOA',
                    19: 'What are the page limitations for FO',
                    20: 'Are there any Limited submission requirements for FOA',
                    21: 'Can I submit more than one application for FOA',
                    22: 'Are Foreign Components allowed for FOA',
                    23: 'What are the other attachments for FOA',
                    25: 'Who are the agency contacts for FOA',
                    26: 'What are the section headers and subsections of the FOA',
                    27: 'What are other submission requirements and information for FOA',
                    28: 'What are the key dates indicated in the table? for FOA',
                    29: 'What are the start dates indicated in the table for FOA'}
    
    main_df = pd.DataFrame({'Questions':dict_questions.values()}, index=dict_questions.keys()).reset_index()
    main_df.columns = ['Question_Number', 'Questions']
    
    return main_df


def create_single_foa_df(gpt_output_dict):
    
    foa_title = list(gpt_output_dict.keys())[0]

    try:
        dict_val = ast.literal_eval(list(gpt_output_dict.values())[0])
        dict_val = {k+1:i for k, i in zip(range(len(dict_val)), dict_val.values())}

    except:
        dict_val = gpt_output_dict.values()
        dict_val = ast.literal_eval(list(dict_val)[0])
        dict_val = {k+1:i for k, i in zip(range(len(dict_val)), dict_val.values())}

    foa_single_df = pd.DataFrame(dict_val.values(), index=dict_val.keys()).reset_index()
    foa_single_df.columns = ['Question_Number', f'{foa_title}']
    
    return foa_single_df

def create_final_df(batch_output_list):
    foa_batch_df = pd.concat([create_single_foa_df(foa) for foa in batch_output_list], axis=1)
    foa_batch_df = foa_batch_df.loc[:,~foa_batch_df.columns.duplicated()]

    final_df = generate_main_df()
    final_df = final_df.merge(foa_batch_df, on=['Question_Number'], how='left')
    
    return final_df

def save_report(final_df, current_datestr):
    save_path = './'
    full_path = os.path.join(save_path, f'foa_reports/{current_datestr}') 
    if not os.path.isdir(full_path):
        os.makedirs(full_path)
    base_fname = f'foa_report_{current_datestr}.xlsx'
    fname = os.path.join(full_path, base_fname)
    print('Saving predictions')
    print(f'File saved as {fname}')
    final_df.to_excel(fname) 
    return fname, base_fname


def create_single_report(doc, foa_name, foa_extract):
    
    par = doc.add_paragraph()
    section_header = f'{foa_name}'
    header = par.add_run(section_header)    
    header.bold = True
    header.font.name = 'Arial'
    header.font.size = docx.shared.Pt(16)
    
    main_text = doc.add_paragraph()
    main_text  = main_text.add_run(foa_extract)
    main_text.font.name = 'Arial'
    main_text.font.size = docx.shared.Pt(12)
    
    doc.add_page_break() 
    
    return doc

def save_final_report(batch_output_dict):
    doc = docx.Document()

    doc = [create_single_report(doc, foa_name, foa_extract) for foa_name, foa_extract in batch_output_dict.items()][0]
    
    datestr = datetime.now().strftime("%Y%m%d_%H%M%S")
    full_path = os.path.join('./', 'reports') 
    if not os.path.isdir(full_path):
        os.makedirs(full_path)
    
    fname = os.path.join(full_path, f'foa_report_{datestr}.docx')
    print(f'Final report saved as {fname}')
    doc.save(fname)
    
    return doc   

def run_analysis(url_list, current_datestr):
    batch_output_dict = {}
    for ix, url in enumerate(url_list):
        print(f'processing {ix}, {url}')
        extracted_text, foa_title = parse_html(url)
        prompt = generate_prompt(extracted_text, current_datestr)
        gpt_output_dict = {foa_title:run_genai(prompt)}
        batch_output_dict.update(gpt_output_dict)
        
    return batch_output_dict


def combine_reports(batch_output_dict):
    final_report = ''
    for foa_name, foa_extract in batch_output_dict.items():
        final_report += f'Report for {foa_name} '
        final_report += '\n'
        final_report += foa_extract 
        final_report += '\n'
        final_report += '\n'
        final_report += '-- END REPORT --'
        final_report += '\n'
        final_report += '\n'
    
    return final_report

def run_pipeline(url_list):
    os.environ['TZ'] = 'US/Eastern'
    time.tzset()
    current_datestr = datetime.now().strftime('%Y%m%d')
    code_version = '0.0.1'
    model_version = '0.0.1'
    
    batch_output_dict = run_analysis(url_list, current_datestr)
    final_report = combine_reports(batch_output_dict)
    save_final_report(batch_output_dict)
    
    return final_report

if __name__ == "__main__":
    final_report = run_pipeline(url_list)